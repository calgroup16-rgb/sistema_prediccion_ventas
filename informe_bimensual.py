import os
import io
import pandas as pd
import matplotlib.pyplot as plt
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

MAPA_BIMESTRES = {
    "B1": {"nombre": "B1 (Ene - Feb)", "texto_titulo": "Enero - Febrero", "meses": [1, 2], "previo": "B6", "meses_prev": [11, 12]},
    "B2": {"nombre": "B2 (Mar - Abr)", "texto_titulo": "Marzo - Abril", "meses": [3, 4], "previo": "B1", "meses_prev": [1, 2]},
    "B3": {"nombre": "B3 (May - Jun)", "texto_titulo": "Mayo - Junio", "meses": [5, 6], "previo": "B2", "meses_prev": [3, 4]},
    "B4": {"nombre": "B4 (Jul - Ago)", "texto_titulo": "Julio - Agosto", "meses": [7, 8], "previo": "B3", "meses_prev": [5, 6]},
    "B5": {"nombre": "B5 (Sep - Oct)", "texto_titulo": "Septiembre - Octubre", "meses": [9, 10], "previo": "B4", "meses_prev": [7, 8]},
    "B6": {"nombre": "B6 (Nov - Dic)", "texto_titulo": "Noviembre - Diciembre", "meses": [11, 12], "previo": "B5", "meses_prev": [9, 10]},
}

def preparar_dataframe(df):
    if df is None or df.empty:
        return df

    df = df.copy()
    if "MES" not in df.columns or "AÑO" not in df.columns:
        col_fecha = None
        for col in df.columns:
            if "FECHA" in str(col).upper():
                col_fecha = col
                break
        
        if col_fecha:
            df[col_fecha] = pd.to_datetime(df[col_fecha], errors='coerce')
            if "MES" not in df.columns:
                df["MES"] = df[col_fecha].dt.month
            if "AÑO" not in df.columns:
                df["AÑO"] = df[col_fecha].dt.year

    return df

def generar_grafica_comparativa(labels, valores_totales):
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    colores = ['#1f77b4', '#ff7f0e', '#2ca02c']
    
    bars = ax.bar(labels, valores_totales, color=colores, width=0.45)
    ax.set_ylabel('Ventas ($)', fontsize=10, fontfamily='sans-serif')
    ax.set_title('Comparativo de Ventas Totales por Período', fontsize=11, fontweight='bold', fontfamily='sans-serif')
    
    for bar in bars:
        height = bar.get_height()
        ax.annotate(f'${height:,.0f}',
                    xy=(bar.get_x() + bar.get_width() / 2, height),
                    xytext=(0, 4),  
                    textcoords="offset points",
                    ha='center', va='bottom', fontsize=9, fontweight='bold', fontfamily='sans-serif')

    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    plt.tight_layout()
    
    img_buf = io.BytesIO()
    plt.savefig(img_buf, format='png', dpi=200)
    img_buf.seek(0)
    plt.close(fig)
    return img_buf

def generar_analisis_y_aspectos(v_act, v_p2, v_p3, df_p1, df_p3, nom_b_act, nom_b_prev, anio_actual, col_val):
    var_vs_anio = ((v_act - v_p2) / v_p2 * 100) if v_p2 > 0 else 0
    var_vs_bim = ((v_act - v_p3) / v_p3 * 100) if v_p3 > 0 else 0

    tend_anio = "un crecimiento" if var_vs_anio >= 0 else "un decrecimiento"
    tend_bim = "un incremento" if var_vs_bim >= 0 else "una disminución"

    # --- ANÁLISIS DE INSUMOS/PRODUCTOS (COLUMNA DESCRIPCION) ---
    col_desc = None
    for c in ["DESCRIPCION", "DESCRIPCIÓN", "PRODUCTO", "CONCEPTO", "LINEA"]:
        if c in df_p1.columns:
            col_desc = c
            break

    texto_insumos = ""
    prod_subieron_str = ""
    prod_bajaron_str = ""

    if col_desc:
        v_p1_prod = df_p1.groupby(col_desc)[col_val].sum() if not df_p1.empty else pd.Series(dtype=float)
        v_p3_prod = df_p3.groupby(col_desc)[col_val].sum() if not df_p3.empty else pd.Series(dtype=float)

        todos_prods = list(set(v_p1_prod.index).union(set(v_p3_prod.index)))
        
        diff_list = []
        for p in todos_prods:
            v_actual = v_p1_prod.get(p, 0.0)
            v_previo = v_p3_prod.get(p, 0.0)
            diff = v_actual - v_previo
            diff_list.append({"producto": p, "diff": diff, "v_act": v_actual, "v_prev": v_previo})

        df_diff = pd.DataFrame(diff_list)

        if not df_diff.empty:
            subieron = df_diff[df_diff["diff"] > 0].sort_values(by="diff", ascending=False).head(5)
            bajaron = df_diff[df_diff["diff"] < 0].sort_values(by="diff", ascending=True).head(5)

            if not subieron.empty:
                prod_subieron_str = ", ".join([f"'{row['producto']}' (+${row['diff']:,.0f})" for _, row in subieron.iterrows()])
            if not bajaron.empty:
                prod_bajaron_str = ", ".join([f"'{row['producto']}' (-${abs(row['diff']):,.0f})" for _, row in bajaron.iterrows()])

            if prod_subieron_str or prod_bajaron_str:
                texto_insumos = " En cuanto al comportamiento de insumos y productos principales: "
                if prod_subieron_str:
                    texto_insumos += f"Se registraron incrementos destacados en las ventas de {prod_subieron_str}."
                if prod_bajaron_str:
                    texto_insumos += f" Por el contrario, se observó una reducción o contracción en productos como {prod_bajaron_str}."

    texto_analisis = (
        f"Durante el período {nom_b_act} de {anio_actual}, se alcanzaron ventas totales netas de clientes de ${v_act:,.2f}. "
        f"Al comparar este resultado con el mismo período del año anterior ({nom_b_act} {anio_actual-1}), "
        f"se evidencia {tend_anio} del {abs(var_vs_anio):.2f}% (frente a ${v_p2:,.2f}). "
        f"Por otra parte, en relación con el comportamiento del bimestre inmediatamente anterior ({nom_b_prev}), se registró "
        f"{tend_bim} del {abs(var_vs_bim):.2f}% respecto a los ${v_p3:,.2f} facturados previamente."
        f"{texto_insumos}"
    )

    col_cli = None
    for c in ["CLIENTE", "NOMBRE_CLIENTE", "TERCERO", "RAZON_SOCIAL"]:
        if c in df_p1.columns:
            col_cli = c
            break

    top_clis = df_p1.groupby(col_cli)[col_val].sum().sort_values(ascending=False).index.tolist() if col_cli and not df_p1.empty else []
    cli_1 = str(top_clis[0]) if len(top_clis) > 0 else "cuentas principales"
    cli_2 = str(top_clis[1]) if len(top_clis) > 1 else "cuentas secundarias"
    cli_3 = str(top_clis[2]) if len(top_clis) > 2 else "otros clientes estratégicos"

    desc_insumos_aspecto = (
        f"Se evidencian variaciones en insumos clave respecto al bimestre anterior. "
        f"Es necesario impulsar la rotación de los insumos que disminuyeron sus ventas "
        f"({prod_bajaron_str if prod_bajaron_str else 'productos con contracción'}) "
        f"y capitalizar la demanda de los insumos con mayor crecimiento ({prod_subieron_str if prod_subieron_str else 'productos en alza'})."
    )

    aspectos = [
        {"titulo": "Recuperación de negocios pendientes y mayor seguimiento comercial", "descripcion": f"Se requiere fortalecer el seguimiento a clientes con procesos de compra pendientes como {cli_1} y {cli_2}, estableciendo fechas de compromiso y realizando un acompañamiento más cercano con las áreas técnicas y de compras."},
        {"titulo": "Análisis y gestión de demanda por líneas e insumos", "descripcion": desc_insumos_aspecto},
        {"titulo": "Mejorar la disponibilidad de inventario en productos estratégicos", "descripcion": "Algunos negocios se vieron afectados por retrasos derivados del desabastecimiento de productos de alta rotación, principalmente sabores, estándares de crioscopio y otros insumos especializados."},
        {"titulo": "Incrementar la participación de las líneas CHARM y productos de mayor rentabilidad", "descripcion": f"Existe una oportunidad importante para fortalecer la comercialización de productos como lactasa, hisopos de luminometria, GMP y demás soluciones CHARM, aprovechando la cartera activa de {cli_3}."},
        {"titulo": "Recuperación de clientes y productos con disminución en ventas", "descripcion": "Se evidencian reducciones en la compra de algunos productos representativos dentro del portafolio. Es necesario identificar las causas de la disminución y presentar alternativas comerciales."},
        {"titulo": "Aumentar la venta cruzada en clientes activos", "descripcion": "Clientes con compras recurrentes representan una oportunidad constante para incrementar la participación mediante la incorporación de nuevas líneas de negocio."},
        {"titulo": "Fortalecer la planeación comercial con clientes estratégicos", "descripcion": "Promover que los clientes compartan sus proyecciones mensuales o trimestrales de consumo permitirá mejorar la planeación de inventarios y anticipar necesidades."},
        {"titulo": "Continuar la recuperación de cuentas con alto potencial", "descripcion": "Se continuará con las gestiones comerciales para recuperar clientes y líneas de negocio que presentan oportunidades de crecimiento."},
        {"titulo": "Optimización de la gestión de cartera y recaudo oportuno", "descripcion": "Monitorear semanalmente la conversión de facturación a cartera efectiva para asegurar el flujo de caja sin frenar la colocación de nuevos pedidos."}
    ]

    texto_cierre = "De acuerdo con la revisión, es importante potenciar la venta de productos como lactasa, hisopos, GMP y en general las líneas de negocio de CHARM."

    return texto_analisis, aspectos, texto_cierre

def dar_formato_tabla(table, col_widths, headers, rows_data):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Encabezado
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)
            
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F4E78')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)

    # Filas de datos
    tot_rows = len(rows_data)
    for r_idx, row in enumerate(rows_data):
        row_cells = table.add_row().cells
        es_ultima_fila = (r_idx == tot_rows - 1)
        
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(9)
                if es_ultima_fila:
                    run.font.bold = True

            if es_ultima_fila:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'D9E1F2')
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)
            elif r_idx % 2 == 1:
                shading = OxmlElement('w:shd')
                shading.set(qn('w:val'), 'clear')
                shading.set(qn('w:color'), 'auto')
                shading.set(qn('w:fill'), 'F2F2F2')
                row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)

def construir_documento_word(contexto, path_template=None):
    if path_template and os.path.exists(path_template):
        doc = Document(path_template)
    else:
        doc = Document()

    # Título Principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run(contexto["titulo_informe"])
    r_title.font.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(31, 78, 120)

    # Subtítulo Responsable
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run(f"Responsable Comercial: {contexto['vendedor']} | Año: {contexto['anio']}")
    r_sub.font.italic = True
    r_sub.font.size = Pt(11)

    doc.add_paragraph()

    # 1. Resumen de Ventas por Fabricante / Proveedor
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. Resumen de Ventas por Fabricante / Proveedor")
    r_h1.font.bold = True
    r_h1.font.size = Pt(13)
    r_h1.font.color.rgb = RGBColor(31, 78, 120)

    headers_prov = ["Fabricante / Proveedor", contexto["col_b_act"], contexto["col_b_ant"], contexto["col_b_prev"]]
    rows_prov = [[item["PROVEEDOR"], item["v_act"], item["v_p2"], item["v_p3"]] for item in contexto["tabla_proveedores"]]
    
    t_prov = doc.add_table(rows=1, cols=4)
    dar_formato_tabla(t_prov, [2.5, 1.3, 1.3, 1.3], headers_prov, rows_prov)

    doc.add_paragraph()

    # 2. Gráfica Comparativa
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(contexto["grafica_ventas"], width=Inches(6.0))

    doc.add_paragraph()

    # 3. Resumen por Cliente
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. Resumen de Ventas por Cliente")
    r_h2.font.bold = True
    r_h2.font.size = Pt(13)
    r_h2.font.color.rgb = RGBColor(31, 78, 120)

    headers_cli = ["Cliente", contexto["col_b_act"], contexto["col_b_ant"], contexto["col_b_prev"]]
    rows_cli = [[item["CLIENTE"], item["v_act"], item["v_p2"], item["v_p3"]] for item in contexto["tabla_clientes"]]
    
    t_cli = doc.add_table(rows=1, cols=4)
    dar_formato_tabla(t_cli, [2.5, 1.3, 1.3, 1.3], headers_cli, rows_cli)

    doc.add_paragraph()

    # 4. Análisis Comercial del Período
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. Análisis Comercial del Período")
    r_h3.font.bold = True
    r_h3.font.size = Pt(13)
    r_h3.font.color.rgb = RGBColor(31, 78, 120)

    p_ana = doc.add_paragraph(contexto["analisis_texto"])
    p_ana.paragraph_format.line_spacing = 1.15

    doc.add_paragraph()

    # 5. Aspectos a Mejorar y Plan de Acción
    h4 = doc.add_paragraph()
    r_h4 = h4.add_run("4. Aspectos a Mejorar y Compromisos Comerciales")
    r_h4.font.bold = True
    r_h4.font.size = Pt(13)
    r_h4.font.color.rgb = RGBColor(31, 78, 120)

    for idx, asp in enumerate(contexto["aspectos_mejorar"], 1):
        p_asp = doc.add_paragraph()
        p_asp.paragraph_format.left_indent = Inches(0.2)
        p_asp.paragraph_format.line_spacing = 1.15
        r_tit = p_asp.add_run(f"• {asp['titulo']}: ")
        r_tit.font.bold = True
        p_asp.add_run(asp["descripcion"])

    doc.add_paragraph()
    p_cierre = doc.add_paragraph(contexto["texto_cierre_aspectos"])
    p_cierre.paragraph_format.line_spacing = 1.15

    return doc

def render_modulo_informe(df_global=None, *args, **kwargs):
    st.title("📄 Generador de Informe Bimensual")

    if df_global is None or (isinstance(df_global, pd.DataFrame) and df_global.empty):
        st.warning("⚠️ Primero debe cargar los datos de ventas en la aplicación.")
        return

    df_global = preparar_dataframe(df_global)

    tab_gen, tab_clientes, tab_fabricantes = st.tabs(["🚀 Generar Informe", "📋 Cargar Lista Clientes", "🏭 Cargar Lista Fabricantes"])

    with tab_clientes:
        st.subheader("Cargar Archivo de Clientes Principales")
        file_cli = st.file_uploader("Cargue la lista de clientes (Excel/CSV):", type=["xlsx", "xls", "csv"], key="file_cli_inf")
        if file_cli:
            try:
                df_c = pd.read_excel(file_cli) if file_cli.name.endswith(('.xlsx', '.xls')) else pd.read_csv(file_cli)
                st.session_state['df_clientes_custom'] = df_c
                st.success("✅ Lista de clientes cargada correctamente.")
                st.dataframe(df_c.head())
            except Exception as e:
                st.error(f"Error al leer archivo de clientes: {e}")

    with tab_fabricantes:
        st.subheader("Cargar Archivo de Fabricantes / Proveedores Principales")
        file_fab = st.file_uploader("Cargue la lista de fabricantes (Excel/CSV):", type=["xlsx", "xls", "csv"], key="file_fab_inf")
        if file_fab:
            try:
                df_f = pd.read_excel(file_fab) if file_fab.name.endswith(('.xlsx', '.xls')) else pd.read_csv(file_fab)
                st.session_state['df_fabricantes_custom'] = df_f
                st.success("✅ Lista de fabricantes cargada correctamente.")
                st.dataframe(df_f.head())
            except Exception as e:
                st.error(f"Error al leer archivo de fabricantes: {e}")

    with tab_gen:
        col_tipo_doc = None
        for col in ["TIPO DOC", "TIPO_DOC", "TIPO DOCUMENTO", "DOCUMENTO", "TIPO"]:
            if col in df_global.columns:
                col_tipo_doc = col
                break

        df_base = df_global.copy()

        col_val = None
        for c in ["TOTAL LINEA", "TOTAL_LINEA", "VALOR_VENTA", "VALOR", "TOTAL"]:
            if c in df_base.columns:
                col_val = c
                break
        
        if not col_val:
            st.error("⚠️ No se encontró la columna 'TOTAL LINEA' en la base de datos.")
            return

        col_vendedor = "VENDEDOR" if "VENDEDOR" in df_base.columns else "RESPONSABLE"
        
        col_cliente = None
        for c in ["CLIENTE", "NOMBRE_CLIENTE", "TERCERO", "RAZON_SOCIAL"]:
            if c in df_base.columns:
                col_cliente = c
                break
                
        col_prov = None
        for c in ["FABRICANTE", "PROVEEDOR", "MARCA"]:
            if c in df_base.columns:
                col_prov = c
                break

        st.subheader("Configuración del Informe")
        col1, col2, col3 = st.columns(3)

        with col1:
            lista_vendedores = list(df_base[col_vendedor].dropna().unique()) if col_vendedor in df_base.columns else ["Todos"]
            vendedor_sel = st.selectbox("Seleccione Vendedor / Responsable:", lista_vendedores)
        with col2:
            bim_sel = st.selectbox("Seleccione el Bimestre del Informe:", list(MAPA_BIMESTRES.keys()), format_func=lambda x: MAPA_BIMESTRES[x]["nombre"])
        with col3:
            anios_disponibles = list(df_base["AÑO"].dropna().astype(int).unique())
            anio_sel = st.selectbox("Seleccione el Año Actual:", anios_disponibles)

        if st.button("🚀 Generar Informe Bimensual"):
            cfg_b = MAPA_BIMESTRES[bim_sel]
            nom_b_act = cfg_b["nombre"]
            nom_b_prev = MAPA_BIMESTRES[cfg_b["previo"]]["nombre"]

            # --- FILTRADO DE CADA PERÍODO (INCLUYENDO FACTURAS Y NOTAS CRÉDITO) ---
            df_p1_base = df_base[(df_base["AÑO"] == anio_sel) & (df_base["MES"].isin(cfg_b["meses"]))]
            if col_tipo_doc:
                df_p1_base = df_p1_base[df_p1_base[col_tipo_doc].astype(str).str.upper().str.contains("FACTURA|NOTA", na=False)]

            df_p2_base = df_base[(df_base["AÑO"] == (anio_sel - 1)) & (df_base["MES"].isin(cfg_b["meses"]))]
            if col_tipo_doc:
                df_p2_base = df_p2_base[df_p2_base[col_tipo_doc].astype(str).str.upper().str.contains("FACTURA|NOTA", na=False)]

            anio_p3 = anio_sel if bim_sel != "B1" else (anio_sel - 1)
            df_p3_base = df_base[(df_base["AÑO"] == anio_p3) & (df_base["MES"].isin(cfg_b["meses_prev"]))]
            if col_tipo_doc:
                df_p3_base = df_p3_base[df_p3_base[col_tipo_doc].astype(str).str.upper().str.contains("FACTURA|NOTA", na=False)]

            # --- OBTENER LISTA DE CLIENTES CARGADOS ---
            lista_cli_custom = []
            if 'df_clientes_custom' in st.session_state and not st.session_state['df_clientes_custom'].empty:
                df_cc = st.session_state['df_clientes_custom']
                col_cc = df_cc.columns[0]
                lista_cli_custom = [str(x).strip().upper() for x in df_cc[col_cc].dropna().unique()]

            if col_cliente and lista_cli_custom:
                df_p1 = df_p1_base[df_p1_base[col_cliente].astype(str).str.strip().str.upper().isin(lista_cli_custom)]
                df_p2 = df_p2_base[df_p2_base[col_cliente].astype(str).str.strip().str.upper().isin(lista_cli_custom)]
                df_p3 = df_p3_base[df_p3_base[col_cliente].astype(str).str.strip().str.upper().isin(lista_cli_custom)]
            else:
                df_p1, df_p2, df_p3 = df_p1_base, df_p2_base, df_p3_base

            # --- 1. TABLA DE CLIENTES (Ordenada de mayor a menor según el período actual) ---
            tabla_clis = []
            if col_cliente and col_cliente in df_base.columns:
                if lista_cli_custom:
                    for cli in lista_cli_custom:
                        va = float(df_p1[df_p1[col_cliente].astype(str).str.strip().str.upper() == cli][col_val].sum()) if not df_p1.empty else 0.0
                        vp2 = float(df_p2[df_p2[col_cliente].astype(str).str.strip().str.upper() == cli][col_val].sum()) if not df_p2.empty else 0.0
                        vp3 = float(df_p3[df_p3[col_cliente].astype(str).str.strip().str.upper() == cli][col_val].sum()) if not df_p3.empty else 0.0

                        tabla_clis.append({"CLIENTE": cli, "v_act_num": va, "v_p2_num": vp2, "v_p3_num": vp3})
                else:
                    clis = df_p1[col_cliente].dropna().unique().tolist()
                    for c in clis:
                        va = float(df_p1[df_p1[col_cliente] == c][col_val].sum()) if not df_p1.empty else 0.0
                        vp2 = float(df_p2[df_p2[col_cliente] == c][col_val].sum()) if not df_p2.empty else 0.0
                        vp3 = float(df_p3[df_p3[col_cliente] == c][col_val].sum()) if not df_p3.empty else 0.0
                        tabla_clis.append({"CLIENTE": str(c), "v_act_num": va, "v_p2_num": vp2, "v_p3_num": vp3})

                # Ordenar clientes de mayor a menor venta en el período actual
                tabla_clis.sort(key=lambda x: x["v_act_num"], reverse=True)

                tot_cli_act = sum(x["v_act_num"] for x in tabla_clis)
                tot_cli_p2 = sum(x["v_p2_num"] for x in tabla_clis)
                tot_cli_p3 = sum(x["v_p3_num"] for x in tabla_clis)

                for item in tabla_clis:
                    item["v_act"] = f"${item['v_act_num']:,.2f}"
                    item["v_p2"] = f"${item['v_p2_num']:,.2f}"
                    item["v_p3"] = f"${item['v_p3_num']:,.2f}"

                tabla_clis.append({
                    "CLIENTE": "TOTAL",
                    "v_act": f"${tot_cli_act:,.2f}",
                    "v_p2": f"${tot_cli_p2:,.2f}",
                    "v_p3": f"${tot_cli_p3:,.2f}"
                })
            else:
                tot_cli_act, tot_cli_p2, tot_cli_p3 = 0.0, 0.0, 0.0

            # --- 2. TABLA DE FABRICANTES (Ordenada de mayor a menor según el período actual) ---
            tabla_provs = []
            if col_prov and col_prov in df_base.columns:
                lista_fab_custom = []
                if 'df_fabricantes_custom' in st.session_state and not st.session_state['df_fabricantes_custom'].empty:
                    df_fc = st.session_state['df_fabricantes_custom']
                    col_fc = df_fc.columns[0]
                    lista_fab_custom = [str(x).strip().upper() for x in df_fc[col_fc].dropna().unique()]

                if lista_fab_custom:
                    otros_v_act, otros_v_p2, otros_v_p3 = 0.0, 0.0, 0.0
                    todos_provs = set(df_p1[col_prov].dropna().astype(str).str.strip()).union(
                        set(df_p2[col_prov].dropna().astype(str).str.strip())
                    ).union(
                        set(df_p3[col_prov].dropna().astype(str).str.strip())
                    )

                    for fab in lista_fab_custom:
                        va = float(df_p1[df_p1[col_prov].astype(str).str.strip().str.upper() == fab][col_val].sum()) if not df_p1.empty else 0.0
                        vp2 = float(df_p2[df_p2[col_prov].astype(str).str.strip().str.upper() == fab][col_val].sum()) if not df_p2.empty else 0.0
                        vp3 = float(df_p3[df_p3[col_prov].astype(str).str.strip().str.upper() == fab][col_val].sum()) if not df_p3.empty else 0.0

                        tabla_provs.append({"PROVEEDOR": fab, "v_act_num": va, "v_p2_num": vp2, "v_p3_num": vp3})

                    # Ordenar fabricantes de la lista de mayor a menor por período actual
                    tabla_provs.sort(key=lambda x: x["v_act_num"], reverse=True)

                    for p in todos_provs:
                        if p.upper() not in lista_fab_custom:
                            otros_v_act += float(df_p1[df_p1[col_prov].astype(str).str.strip() == p][col_val].sum()) if not df_p1.empty else 0.0
                            otros_v_p2 += float(df_p2[df_p2[col_prov].astype(str).str.strip() == p][col_val].sum()) if not df_p2.empty else 0.0
                            otros_v_p3 += float(df_p3[df_p3[col_prov].astype(str).str.strip() == p][col_val].sum()) if not df_p3.empty else 0.0

                    tabla_provs.append({"PROVEEDOR": "Otros", "v_act_num": otros_v_act, "v_p2_num": otros_v_p2, "v_p3_num": otros_v_p3})
                else:
                    provs = df_p1[col_prov].dropna().unique().tolist()
                    for p in provs:
                        va = float(df_p1[df_p1[col_prov] == p][col_val].sum()) if not df_p1.empty else 0.0
                        vp2 = float(df_p2[df_p2[col_prov] == p][col_val].sum()) if not df_p2.empty else 0.0
                        vp3 = float(df_p3[df_p3[col_prov] == p][col_val].sum()) if not df_p3.empty else 0.0
                        tabla_provs.append({"PROVEEDOR": str(p), "v_act_num": va, "v_p2_num": vp2, "v_p3_num": vp3})

                    # Ordenar fabricantes de mayor a menor por período actual
                    tabla_provs.sort(key=lambda x: x["v_act_num"], reverse=True)

                tot_fab_act = sum(x["v_act_num"] for x in tabla_provs)
                tot_fab_p2 = sum(x["v_p2_num"] for x in tabla_provs)
                tot_fab_p3 = sum(x["v_p3_num"] for x in tabla_provs)

                for item in tabla_provs:
                    item["v_act"] = f"${item['v_act_num']:,.2f}"
                    item["v_p2"] = f"${item['v_p2_num']:,.2f}"
                    item["v_p3"] = f"${item['v_p3_num']:,.2f}"

                tabla_provs.append({
                    "PROVEEDOR": "TOTAL",
                    "v_act": f"${tot_fab_act:,.2f}",
                    "v_p2": f"${tot_fab_p2:,.2f}",
                    "v_p3": f"${tot_fab_p3:,.2f}"
                })

            head_b_act = f"VENTA {bim_sel} {anio_sel}"
            head_b_ant_anio = f"VENTA {bim_sel} {anio_sel-1}"
            head_b_prev = f"VENTA {cfg_b['previo']} {anio_p3}"

            buf_grafica = generar_grafica_comparativa([head_b_act, head_b_ant_anio, head_b_prev], [tot_cli_act, tot_cli_p2, tot_cli_p3])
            txt_analisis, aspectos_lista, txt_cierre = generar_analisis_y_aspectos(tot_cli_act, tot_cli_p2, tot_cli_p3, df_p1, df_p3, nom_b_act, nom_b_prev, anio_sel, col_val)

            path_template = "templates/BIMENSUAL MAY-JUN.docx"

            contexto = {
                "vendedor": str(vendedor_sel),
                "periodo_titulo": str(cfg_b["texto_titulo"]),
                "titulo_informe": f"INFORME BIMENSUAL {cfg_b['texto_titulo'].upper()}",
                "anio": str(anio_sel),
                "col_b_act": head_b_act,
                "col_b_ant": head_b_ant_anio,
                "col_b_prev": head_b_prev,
                "tabla_proveedores": tabla_provs,
                "tabla_clientes": tabla_clis,
                "analisis_texto": txt_analisis,
                "aspectos_mejorar": aspectos_lista,
                "texto_cierre_aspectos": txt_cierre,
                "grafica_ventas": buf_grafica
            }

            try:
                doc = construir_documento_word(contexto, path_template=path_template)
                
                output_buf = io.BytesIO()
                doc.save(output_buf)
                output_buf.seek(0)

                st.success("✅ Informe generado correctamente.")
                st.download_button(
                    label="📥 Descargar Informe Word (.docx)",
                    data=output_buf,
                    file_name=f"INFORME_BIMENSUAL_{bim_sel}_{vendedor_sel}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as err:
                st.error(f"⚠️ Error al construir el documento Word: {err}")
